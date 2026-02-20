#!/usr/bin/env python3
"""Create meeting_report.docx template for Foxit Document Generation API.
Uses python-docx when available for a fully valid Word document; falls back to zipfile."""
import os
import zipfile
from io import BytesIO

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def create_with_docx(out_path: str) -> bool:
    """Create template using python-docx (produces valid .docx)."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Meeting Report", 0)
        doc.add_paragraph("Title: {{title}}")
        doc.add_paragraph("Date: {{today}}  |  Duration: {{duration}}  |  Words: {{wordCount}}")
        doc.add_paragraph()
        doc.add_heading("Summary", level=1)
        doc.add_paragraph("{{summary}}")
        doc.add_paragraph()
        doc.add_heading("Key Insights", level=1)
        doc.add_paragraph("{{keyInsightsText}}")
        doc.add_paragraph()
        doc.add_heading("Decisions", level=1)
        doc.add_paragraph("{{decisionsText}}")
        doc.add_paragraph()
        doc.add_heading("Action Items", level=1)
        doc.add_paragraph("{{actionItemsText}}")
        doc.add_paragraph()
        doc.add_heading("Full Transcript", level=1)
        doc.add_paragraph("{{fullTranscript}}")
        doc.save(out_path)
        return True
    except ImportError:
        return False


def create_with_zipfile(out_path: str) -> None:
    """Fallback: create minimal .docx using zipfile (no python-docx)."""
    def para(text: str) -> str:
        return f'<w:p xmlns:w="{W_NS}"><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'

    doc_body = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="{W_NS}"><w:body>
{para("Meeting Report")}
{para("Title: {{title}}")}
{para("Date: {{today}}  |  Duration: {{duration}}  |  Words: {{wordCount}}")}
{para("")}
{para("Summary")}
{para("{{summary}}")}
{para("")}
{para("Key Insights")}
{para("{{keyInsightsText}}")}
{para("")}
{para("Decisions")}
{para("{{decisionsText}}")}
{para("")}
{para("Action Items")}
{para("{{actionItemsText}}")}
{para("")}
{para("Full Transcript")}
{para("{{fullTranscript}}")}
<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>
</w:body></w:document>'''

    content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>'''

    rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>'''

    doc_rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>
'''
    core = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties">
<dc:title xmlns:dc="http://purl.org/dc/elements/1.1/">Meeting Report</dc:title>
</cp:coreProperties>'''
    app = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>Insightly</Application></Properties>'''

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", doc_body)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("docProps/core.xml", core)
        zf.writestr("docProps/app.xml", app)
    with open(out_path, "wb") as f:
        f.write(buf.getvalue())


def main():
    out_dir = os.path.join(os.path.dirname(__file__), "..", "templates")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "meeting_report.docx")
    if create_with_docx(out_path):
        print(f"Created {out_path} (python-docx)")
    else:
        create_with_zipfile(out_path)
        print(f"Created {out_path} (zipfile fallback)")


if __name__ == "__main__":
    main()
